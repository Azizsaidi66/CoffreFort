from fastapi import FastAPI, Depends, HTTPException, Form, UploadFile, File, BackgroundTasks, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy import create_engine, Column, String, DateTime, Boolean, Integer, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from datetime import datetime, timedelta
from typing import Optional, List
import os
import requests
import json
import uuid
from passlib.context import CryptContext
from jose import JWTError, jwt
import httpx

# ==================== CONFIG ====================
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://mayan:mayanpassword@postgres:5432/coffre_fort")
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-change-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = int(os.getenv("ACCESS_TOKEN_EXPIRE_MINUTES", "30"))
OLLAMA_API_URL = os.getenv("OLLAMA_API_URL", "http://ollama:11434")
MAYAN_API_URL = os.getenv("MAYAN_API_URL", "http://mayan:8000")
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/app/uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ==================== DATABASE ====================
engine = create_engine(DATABASE_URL, future=True)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# ==================== MODELS ====================
class User(Base):
    __tablename__ = "users"
    
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String, unique=True, index=True)
    hashed_password = Column(String)
    full_name = Column(String)
    role = Column(String, default="user")  # "user" or "admin"
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=datetime.utcnow)

class AccessWindow(Base):
    __tablename__ = "access_windows"
    
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, index=True)
    start_time = Column(String)  # "HH:MM"
    end_time = Column(String)    # "HH:MM"
    created_at = Column(DateTime, default=datetime.utcnow)

class Document(Base):
    __tablename__ = "documents"
    
    id = Column(Integer, primary_key=True, index=True)
    mayan_id = Column(String, unique=True, index=True)
    title = Column(String)
    description = Column(Text)
    uploaded_by = Column(Integer, index=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    ai_summary = Column(Text, nullable=True)
    ai_keywords = Column(Text, nullable=True)
    file_path = Column(String, nullable=True)

class SessionModel(Base):
    __tablename__ = "sessions"
    
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, index=True)
    token = Column(String, unique=True, index=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    expires_at = Column(DateTime)

# Create tables
Base.metadata.create_all(bind=engine)

# ==================== SECURITY ====================
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def hash_password(password: str) -> str:
    if not isinstance(password, str):
        password = str(password)
    pw_bytes = password.encode("utf-8")
    if len(pw_bytes) > 72:
        # reject or truncate; here we reject at API level, but keep safe truncation for hashing
        password = pw_bytes[:72].decode("utf-8", errors="ignore")
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str) -> bool:
    if not isinstance(plain_password, str):
        plain_password = str(plain_password)
    pw_bytes = plain_password.encode("utf-8")
    if len(pw_bytes) > 72:
        plain_password = pw_bytes[:72].decode("utf-8", errors="ignore")
    return pwd_context.verify(plain_password, hashed_password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

# ==================== DEPENDENCIES ====================
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

async def get_current_user(authorization: Optional[str] = Header(None), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    if not authorization:
        raise credentials_exception
    parts = authorization.split()
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise credentials_exception
    token = parts[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: int = payload.get("sub")
        if user_id is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = db.query(User).filter(User.id == user_id).first()
    if user is None:
        raise credentials_exception
    return user

# ==================== FASTAPI APP ====================
app = FastAPI(title="CoffreFort Backend", version="1.0.0")
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== AUTH ENDPOINTS ====================
@app.post("/auth/register")
async def register(
    email: str = Form(...),
    password: str = Form(...),
    full_name: str = Form(...),
    db: Session = Depends(get_db)
):
    existing_user = db.query(User).filter(User.email == email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    if len(password.encode("utf-8")) > 72:
        raise HTTPException(status_code=400, detail="Password too long (max 72 bytes)")

    user = User(
        email=email,
        hashed_password=hash_password(password),
        full_name=full_name,
        role="user"
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    
    access_token = create_access_token(data={"sub": user.id})
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user_id": user.id,
        "role": user.role
    }

@app.post("/auth/login")
async def login(
    email: str = Form(...),
    password: str = Form(...),
    role: str = Form(...),
    db: Session = Depends(get_db)
):
    """User login"""
    user = db.query(User).filter(User.email == email).first()
    if not user or not verify_password(password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if user.role != role:
        raise HTTPException(status_code=403, detail="Role mismatch")
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.id},
        expires_delta=access_token_expires
    )
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user_id": user.id,
        "role": user.role,
        "email": user.email
    }

# ==================== USER MANAGEMENT ====================
@app.get("/users/me")
async def get_current_user_info(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    return {
        "id": current_user.id,
        "email": current_user.email,
        "full_name": current_user.full_name,
        "role": current_user.role,
        "created_at": current_user.created_at
    }

@app.get("/users")
async def list_users(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    users = db.query(User).all()
    return [
        {
            "id": u.id,
            "email": u.email,
            "full_name": u.full_name,
            "role": u.role,
            "is_active": u.is_active
        }
        for u in users
    ]

@app.post("/users")
async def create_user(
    email: str = Form(...),
    password: str = Form(...),
    full_name: str = Form(...),
    role: str = Form("user"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create new user (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    existing_user = db.query(User).filter(User.email == email).first()
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already exists")
    
    if len(password.encode("utf-8")) > 72:
        raise HTTPException(status_code=400, detail="Password too long (max 72 bytes).")

    user = User(
        email=email,
        hashed_password=hash_password(password),
        full_name=full_name,
        role=role
    )
    db.add(user)
    db.commit()
    db.refresh(user)
    
    return {"id": user.id, "email": user.email, "role": user.role}

@app.delete("/users/{user_id}")
async def delete_user(
    user_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    db.delete(user)
    db.commit()
    return {"message": "User deleted"}

# ==================== ACCESS WINDOWS ====================
@app.post("/access-windows")
async def set_access_window(
    user_id: int,
    start_time: str,
    end_time: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Delete existing window
    db.query(AccessWindow).filter(AccessWindow.user_id == user_id).delete()
    
    window = AccessWindow(
        user_id=user_id,
        start_time=start_time,
        end_time=end_time
    )
    db.add(window)
    db.commit()
    
    return {"message": "Access window updated"}

@app.get("/access-windows/{user_id}")
async def get_access_window(
    user_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    window = db.query(AccessWindow).filter(AccessWindow.user_id == user_id).first()
    if not window:
        return {"start_time": "00:00", "end_time": "23:59"}
    
    return {"start_time": window.start_time, "end_time": window.end_time}

def check_access_allowed(current_time: str, window_start: str, window_end: str) -> bool:
    current = datetime.strptime(current_time, "%H:%M").time()
    start = datetime.strptime(window_start, "%H:%M").time()
    end = datetime.strptime(window_end, "%H:%M").time()
    
    if start <= end:
        return start <= current <= end
    else:  # Window crosses midnight
        return current >= start or current <= end

@app.get("/check-access")
async def check_access(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    window = db.query(AccessWindow).filter(AccessWindow.user_id == current_user.id).first()
    
    if not window:
        return {"allowed": True}
    
    current_time = datetime.now().strftime("%H:%M")
    allowed = check_access_allowed(current_time, window.start_time, window.end_time)
    
    return {
        "allowed": allowed,
        "current_time": current_time,
        "window_start": window.start_time,
        "window_end": window.end_time
    }

# ==================== AI ANALYSIS ====================
async def analyze_with_ollama(text: str, background_tasks: BackgroundTasks) -> dict:
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            # Pull model if needed
            await client.post(f"{OLLAMA_API_URL}/api/pull", json={"name": "mistral"})
            
            # Generate summary
            response = await client.post(
                f"{OLLAMA_API_URL}/api/generate",
                json={
                    "model": "mistral",
                    "prompt": f"Summarize the following document in French in 3-4 sentences:\n\n{text[:2000]}",
                    "stream": False
                }
            )
            
            if response.status_code == 200:
                summary = response.json().get("response", "")
                
                keyword_response = await client.post(
                    f"{OLLAMA_API_URL}/api/generate",
                    json={
                        "model": "mistral",
                        "prompt": f"Extract 8-10 key words/phrases from this text in French, separated by commas:\n\n{text[:2000]}",
                        "stream": False
                    }
                )
                
                keywords = keyword_response.json().get("response", "")
                
                return {
                    "summary": summary.strip(),
                    "keywords": keywords.strip()
                }
    except Exception as e:
        print(f"Ollama error: {e}")
    
    return {"summary": "Analysis failed", "keywords": ""}

@app.post("/documents/analyze")
async def analyze_document(
    document_id: int,
    text: str = Form(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # permission
    if current_user.role != "admin" and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403, detail="Not allowed to analyze this document")
    
    analysis = await analyze_with_ollama(text, background_tasks)
    
    doc.ai_summary = analysis["summary"]
    doc.ai_keywords = analysis["keywords"]
    db.commit()
    
    return analysis

# helper: extract text from uploaded file
async def extract_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    # PDF
    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            return "\n".join(text_parts)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF parsing failed: {e}")
    # DOCX
    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX parsing failed: {e}")
    # Plain text fallback
    try:
        with open(file_path, "rb") as f:
            raw = f.read()
            return raw.decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File read failed: {e}")

# analyze by document id using the uploaded file
@app.post("/documents/analyze-file/{document_id}")
async def analyze_file_from_upload(
    document_id: int,
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if current_user.role != "admin" and doc.uploaded_by != current_user.id:
        raise HTTPException(status_code=403, detail="Not allowed")
    if not getattr(doc, "file_path", None):
        raise HTTPException(status_code=400, detail="No file attached")
    file_path = os.path.join(UPLOAD_DIR, doc.file_path)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Uploaded file not found")

    text = await extract_text_from_file(file_path)
    analysis = await analyze_with_ollama(text, background_tasks)

    doc.ai_summary = analysis.get("summary", "")
    doc.ai_keywords = analysis.get("keywords", "")
    db.commit()
    db.refresh(doc)

    return {
        "document_id": doc.id,
        "summary": doc.ai_summary,
        "keywords": doc.ai_keywords,
        "file_url": f"/uploads/{doc.file_path}"
    }

# ==================== DOCUMENT MANAGEMENT ====================
@app.post("/documents/upload")
async def upload_document(
    title: str = Form(...),
    description: str = Form(""),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Allow any authenticated user to upload; record uploader id."""
    try:
        ext = os.path.splitext(file.filename)[1]
        unique_name = f"{uuid.uuid4().hex}{ext}"
        dest_path = os.path.join(UPLOAD_DIR, unique_name)
        content = await file.read()
        with open(dest_path, "wb") as f:
            f.write(content)
        
        document = Document(
            mayan_id=f"doc_{current_user.id}_{datetime.utcnow().timestamp()}",
            title=title,
            description=description,
            uploaded_by=current_user.id,
            file_path=unique_name
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        return {
            "id": document.id,
            "mayan_id": document.mayan_id,
            "title": title,
            "created_at": document.created_at,
            "file_url": f"/uploads/{unique_name}"
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/documents")
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    if current_user.role == "admin":
        documents = db.query(Document).all()
    else:
        documents = db.query(Document).filter(Document.uploaded_by == current_user.id).all()

    return [
        {
            "id": d.id,
            "title": d.title,
            "description": d.description,
            "uploaded_by": d.uploaded_by,
            "created_at": d.created_at,
            "ai_summary": d.ai_summary,
            "ai_keywords": d.ai_keywords,
            "file_url": (f"/uploads/{d.file_path}" if getattr(d, "file_path", None) else None)
        }
        for d in documents
    ]

@app.get("/documents/{document_id}")
async def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": doc.id,
        "title": doc.title,
        "description": doc.description,
        "ai_summary": doc.ai_summary,
        "ai_keywords": doc.ai_keywords,
        "created_at": doc.created_at,
        "file_url": (f"/uploads/{doc.file_path}" if getattr(doc, "file_path", None) else None)
    }



# ==================== MAYAN INTEGRATION ====================
@app.post("/mayan/sso-token")
async def generate_mayan_sso_token(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    sso_token = create_access_token(
        data={
            "sub": current_user.id,
            "email": current_user.email,
            "mayan_user": current_user.email
        },
        expires_delta=timedelta(hours=1)
    )
    
    return {
        "sso_token": sso_token,
        "mayan_url": f"{MAYAN_API_URL}?token={sso_token}"
    }


@app.get("/")
async def root():
    return {
        "message": "CoffreFort Backend is running!",
        "version": "1.0.0",
        "docs": "/docs"
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8000")))